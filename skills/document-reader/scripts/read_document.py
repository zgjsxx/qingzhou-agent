from __future__ import annotations

import argparse
import csv
import json
from pathlib import Path


TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".log", ".json", ".yaml", ".yml", ".xml"}
WORD_EXTENSIONS = {".docx"}
DELIMITED_EXTENSIONS = {".csv", ".tsv"}


def truncate(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    return f"{text[:max_chars]}\n... ({len(text) - max_chars} more characters)"


def format_rows(rows: list[list[str]], max_cols: int) -> list[str]:
    lines: list[str] = []
    for row in rows:
        cells = [str(cell or "").replace("\n", " ").strip() for cell in row[:max_cols]]
        if len(row) > max_cols:
            cells.append(f"... ({len(row) - max_cols} more columns)")
        lines.append(" | ".join(cells))
    return lines


def read_text(path: Path, max_chars: int) -> str:
    return truncate(path.read_text(encoding="utf-8", errors="replace"), max_chars)


def read_docx(path: Path, max_chars: int, include_tables: bool) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise SystemExit("python-docx is not installed. Run: pip install -r requirements.txt") from exc

    document = Document(str(path))
    lines = [
        "[Word document]",
        f"path: {path}",
        f"paragraphs: {len(document.paragraphs)}",
        f"tables: {len(document.tables)}",
        "",
        "Text:",
    ]
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    if include_tables and document.tables:
        lines.extend(["", "Tables:"])
        for table_index, table in enumerate(document.tables, start=1):
            lines.append(f"[table {table_index}]")
            for row in table.rows:
                lines.append(" | ".join(cell.text.strip() for cell in row.cells))

    return truncate("\n".join(lines), max_chars)


def read_delimited(path: Path, max_chars: int, max_rows: int, max_cols: int) -> str:
    delimiter = "\t" if path.suffix.lower() == ".tsv" else ","
    rows: list[list[str]] = []
    row_count = 0
    with path.open("r", encoding="utf-8-sig", errors="replace", newline="") as handle:
        reader = csv.reader(handle, delimiter=delimiter)
        for row in reader:
            row_count += 1
            if row_count <= max_rows:
                rows.append(row)

    lines = [
        "[Delimited document]",
        f"path: {path}",
        f"type: {path.suffix.lower().lstrip('.')}",
        f"rows: {row_count}",
        "",
        "[preview]",
        *format_rows(rows, max_cols),
    ]
    if row_count > max_rows:
        lines.append(f"... ({row_count - max_rows} more rows)")
    return truncate("\n".join(lines), max_chars)


def main() -> int:
    parser = argparse.ArgumentParser(description="Read text, CSV/TSV, and DOCX documents.")
    parser.add_argument("path", help="Path to the document.")
    parser.add_argument("--max-chars", type=int, default=12000)
    parser.add_argument("--max-rows", type=int, default=30)
    parser.add_argument("--max-cols", type=int, default=12)
    parser.add_argument("--no-tables", action="store_true", help="Skip DOCX tables.")
    parser.add_argument("--json", action="store_true", help="Emit JSON with metadata and text.")
    args = parser.parse_args()

    path = Path(args.path).expanduser().resolve()
    if not path.exists() or not path.is_file():
        raise SystemExit(f"file does not exist: {path}")

    suffix = path.suffix.lower()
    if suffix in TEXT_EXTENSIONS:
        text = read_text(path, args.max_chars)
        kind = "text"
    elif suffix in WORD_EXTENSIONS:
        text = read_docx(path, args.max_chars, include_tables=not args.no_tables)
        kind = "word"
    elif suffix in DELIMITED_EXTENSIONS:
        text = read_delimited(path, args.max_chars, args.max_rows, args.max_cols)
        kind = "delimited"
    elif suffix == ".doc":
        raise SystemExit("legacy .doc files are not directly supported; convert to .docx first")
    else:
        raise SystemExit(f"unsupported document type: {suffix or '(none)'}")

    if args.json:
        print(json.dumps({"path": str(path), "kind": kind, "text": text}, ensure_ascii=False, indent=2))
    else:
        print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
