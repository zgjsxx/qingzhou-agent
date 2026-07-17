import subprocess
import sys
import tempfile
import unittest
from pathlib import Path

from tools import registry


ROOT = Path(__file__).resolve().parents[1]
DOCUMENT_READER = ROOT / "skills" / "document-reader" / "scripts" / "read_document.py"
SPREADSHEET_READER = ROOT / "skills" / "spreadsheet-reader" / "scripts" / "read_spreadsheet.py"


class DocumentReaderScriptsTest(unittest.TestCase):
    def run_script(self, script: Path, *args: str) -> str:
        result = subprocess.run(
            [sys.executable, str(script), *args],
            cwd=ROOT,
            text=True,
            encoding="utf-8",
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True,
        )
        return result.stdout

    def test_document_reader_extracts_docx_paragraphs_and_tables(self):
        try:
            from docx import Document
        except ImportError:  # pragma: no cover - dependency is in requirements.
            self.skipTest("python-docx is not installed")

        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "sample.docx"
            document = Document()
            document.add_paragraph("Project summary")
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Name"
            table.cell(0, 1).text = "Value"
            table.cell(1, 0).text = "Status"
            table.cell(1, 1).text = "Ready"
            document.save(path)

            output = self.run_script(DOCUMENT_READER, str(path))

        self.assertIn("[Word document]", output)
        self.assertIn("Project summary", output)
        self.assertIn("Name | Value", output)
        self.assertIn("Status | Ready", output)

    def test_spreadsheet_reader_reads_csv_preview(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "data.csv"
            path.write_text("name,value\nalpha,1\nbeta,2\n", encoding="utf-8")

            output = self.run_script(
                SPREADSHEET_READER,
                str(path),
                "--max-rows",
                "2",
                "--max-cols",
                "2",
            )

        self.assertIn("[Spreadsheet]", output)
        self.assertIn("type: csv", output)
        self.assertIn("name | value", output)
        self.assertIn("alpha | 1", output)
        self.assertIn("... (1 more rows)", output)

    def test_document_reader_reads_text_files(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            path = Path(tmpdir) / "notes.txt"
            path.write_text("hello notes", encoding="utf-8")

            output = self.run_script(DOCUMENT_READER, str(path))

        self.assertEqual(output.strip(), "hello notes")

    def test_reader_capabilities_are_not_registered_as_agent_tools(self):
        names = {getattr(tool, "name", "") for tool in registry.ALL_TOOLS}

        self.assertNotIn("read_uploaded_document", names)
        self.assertNotIn("read_word_document", names)
        self.assertNotIn("read_spreadsheet", names)


if __name__ == "__main__":
    unittest.main()
